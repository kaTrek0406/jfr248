import asyncio
import glob
import os
import re
import sqlite3
import tempfile
import zipfile
from datetime import datetime, timedelta, date

# ===== Timezone =====
try:
    from zoneinfo import ZoneInfo
except Exception:
    ZoneInfo = None

from aiogram import Bot, Dispatcher, F, Router
from aiogram.filters import Command
from aiogram.types import Message, CallbackQuery, ReplyKeyboardMarkup, KeyboardButton
from aiogram.utils.keyboard import InlineKeyboardBuilder
from aiogram.client.default import DefaultBotProperties
from aiogram.exceptions import TelegramBadRequest
from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

# ================== ENV / TZ ==================
load_dotenv()
BOT_TOKEN = os.getenv("BOT_TOKEN", "")
DB_PATH = os.getenv("DB_PATH", "./schedule.db")
TZNAME = os.getenv("TZ", "Europe/Chisinau")
DEFAULT_GROUP = os.getenv("DEFAULT_GROUP", "JFR-237")
DOCX_PATH = os.getenv("DOCX_PATH", "").strip()
DOCX_GLOB = os.getenv("DOCX_GLOB", "").strip()

ROMAN_PREFIX = re.compile(r"^\s*(?:I|II|III|IV|V|VI|VII|VIII|IX|X)\s+", re.IGNORECASE)

def parse_ora_cell(ora: str):
    o = normalize_sup(ora)
    o = ROMAN_PREFIX.sub("", o)
    m = re.search(r'(\d{1,2})[:\. ]?(\d{2})\s*-\s*(\d{1,2})[:\. ]?(\d{2})', o)
    if not m:
        return None, None
    h1, m1, h2, m2 = map(int, m.groups())
    return f"{h1:02d}:{m1:02d}", f"{h2:02d}:{m2:02d}"


def get_tz():
    if ZoneInfo is not None:
        try:
            return ZoneInfo(TZNAME)
        except Exception:
            pass
    from datetime import timezone
    return timezone(timedelta(hours=3))  # Chisinau UTC+3 (fallback, без DST)
TZ = get_tz()

# ================== GROUPS ==================
GROUP_LABELS = {
    "JFR-237": "Jurnalism și procese mediatice — JFR-237",
}
GROUP_ALIASES = {
    "Jurnalism și procese mediatice": "JFR-237",
    "JFR - 237": "JFR-237",
    "JFR-237": "JFR-237",
}

# ================== DB ==================
SCHEMA = """
CREATE TABLE IF NOT EXISTS events (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    date TEXT NOT NULL,
    time_start TEXT,
    time_end TEXT,
    title TEXT NOT NULL,
    teacher TEXT,
    room TEXT,
    group_code TEXT NOT NULL
);
CREATE INDEX IF NOT EXISTS idx_events_date_group ON events(date, group_code);
"""



def db():
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    return con

def init_db():
    with db() as con:
        for stmt in SCHEMA.strip().split(";"):
            s = stmt.strip()
            if s:
                con.execute(s + ";")

def clear_group(group_code: str):
    with db() as con:
        con.execute("DELETE FROM events WHERE group_code=?", (group_code,))

def insert_event(date_str, t_start, t_end, title, teacher, room, group_code):
    with db() as con:
        con.execute(
            "INSERT INTO events(date, time_start, time_end, title, teacher, room, group_code) VALUES(?,?,?,?,?,?,?)",
            (date_str, t_start, t_end, title, teacher, room, group_code),
        )

def fetch_day_for_group(d: date, group_code: str):
    ymd = d.strftime("%Y-%m-%d")
    with db() as con:
        cur = con.execute(
            """SELECT * FROM events
               WHERE date=? AND group_code=?
               ORDER BY COALESCE(time_start,'99:99'), title""",
            (ymd, group_code),
        )
        return cur.fetchall()

def fetch_week_for_group(start: date, end: date, group_code: str):
    with db() as con:
        cur = con.execute(
            """SELECT * FROM events
               WHERE date BETWEEN ? AND ? AND group_code=?
               ORDER BY date, COALESCE(time_start,'99:99'), title""",
            (start.strftime("%Y-%m-%d"), end.strftime("%Y-%m-%d"), group_code),
        )
        return cur.fetchall()

# ================== Helpers / UI ==================
def week_bounds(any_day: date):
    start = any_day - timedelta(days=any_day.weekday())  # Monday
    end = start + timedelta(days=6)
    return start, end

def fmt_pair(row):
    t1 = row["time_start"] or ""
    t2 = row["time_end"] or ""
    time_part = f"{t1}-{t2}" if (t1 and t2) else (t1 or "—")
    title = row["title"] or ""
    teacher = row["teacher"] or ""
    room = row["room"] or ""
    line = f"• <b>{time_part}</b> · <i>{title}</i>"
    extras = []
    if teacher: extras.append(f"👩‍🏫 {teacher}")
    if room: extras.append(f"🏫 <b>{room}</b>")
    if extras:
        line += "\n  " + " · ".join(extras)
    return line

def fmt_day_block(d: date, rows, group_code: str):
    head = f"📅 {d.strftime('%a, %d.%m.%Y')} — <b>{GROUP_LABELS.get(group_code, group_code)}</b>"
    if not rows:
        return f"{head}\nЗанятий нет."
    items = "\n".join(fmt_pair(r) for r in rows)
    return f"{head}\n{items}"

def main_menu():
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="🗓 Выбрать день недели"), KeyboardButton(text="📅 Сегодня")],
            [KeyboardButton(text="👥 Выбрать группу"), KeyboardButton(text="🗂 Эта неделя")],
        ],
        resize_keyboard=True,
    )

def days_keyboard(anchor: date):
    start, end = week_bounds(anchor)
    b = InlineKeyboardBuilder()

    prev_txt = (start - timedelta(days=7)).strftime("%d.%m")
    next_txt = (start + timedelta(days=7)).strftime("%d.%m")
    mid_txt = f"Неделя {start.strftime('%d.%m')}–{end.strftime('%d.%m')}"

    b.button(text=f"◀️ {prev_txt}", callback_data=f"wk:{(start - timedelta(days=7)).isoformat()}")
    b.button(text=mid_txt, callback_data="noop")
    b.button(text=f"▶️ {next_txt}", callback_data=f"wk:{(start + timedelta(days=7)).isoformat()}")
    b.adjust(3)

    weekdays = ["Пн","Вт","Ср","Чт","Пт","Сб","Вс"]
    d = start
    for i in range(7):
        label = f"{weekdays[i]} {d.day:02d}"
        if d == datetime.now(TZ).date():
            label = f"🟩 {label}"
        b.button(text=label, callback_data=f"d:{d.isoformat()}")
        d += timedelta(days=1)
    b.adjust(3,4)
    return b.as_markup()

def groups_kb():
    b = InlineKeyboardBuilder()
    for g in ["JFR-237"]:
        b.button(text=g, callback_data=f"g:{g}")
    b.adjust(1)
    return b.as_markup()

async def safe_edit(message, *, text=None, reply_markup=None):
    """Безопасное редактирование (тихо игнорирует 'message is not modified')."""
    try:
        if text is not None:
            await message.edit_text(text, reply_markup=reply_markup)
        else:
            await message.edit_reply_markup(reply_markup=reply_markup)
    except TelegramBadRequest as e:
        if "message is not modified" in str(e).lower():
            if text is not None and reply_markup is not None:
                try:
                    await message.edit_reply_markup(reply_markup=reply_markup)
                except TelegramBadRequest as e2:
                    if "message is not modified" not in str(e2).lower():
                        raise
        else:
            raise

# ================== Parsing helpers ==================
SUP = str.maketrans({
    "⁰":"0","¹":"1","²":"2","³":"3","⁴":"4","⁵":"5","⁶":"6","⁷":"7","⁸":"8","⁹":"9",
    "º":"0","˙":"", "’":"'", "ː":":"
})

def normalize_sup(s: str) -> str:
    s = (s or "").translate(SUP)
    s = s.replace("–","-").replace("—","-").replace("‒","-")
    return re.sub(r"\s+", " ", s).strip()
def parse_ora_cell(ora: str):
    text = ROMAN_PREFIX.sub("", normalize_sup(ora))
    m = re.search(r'(\d{1,2})[:\. ]?(\d{2})\s*-\s*(\d{1,2})[:\. ]?(\d{2})', text)
    if not m: return None, None
    h1, m1, h2, m2 = map(int, m.groups())
    return f"{h1:02d}:{m1:02d}", f"{h2:02d}:{m2:02d}"


def parse_group_cell(txt: str):
    t = (txt or "").strip()
    if not t: return None, None, None
    lines = [re.sub(r"\s+", " ", x.strip()) for x in t.splitlines() if x.strip()]
    if not lines: return None, None, None

    title = lines[0]
    teacher = None
    room = None

    # аудитория: допускаем 423/biblioteca, Bl.2/s.102 и т.п.
    for line in lines[1:]:
        m = re.search(r'\b(?:sală|s\.?|aud\.?|cab\.?|ауд\.?)\s*([0-9A-Za-z./\\\- ]+)', line, re.IGNORECASE)
        if m and not room:
            room = m.group(1).strip()

    for line in lines[1:]:
        if room and re.search(r'\b(?:sală|s\.?|aud\.?|cab\.?|ауд\.?)\b', line, re.IGNORECASE):
            continue
        if any(w in line.lower() for w in ["dr.", "conf.", "prof.", "lector", "asistent", "universitar", "univ", "cadru"]):
            teacher = line; break
        if re.search(r'[A-ZĂÂÎȘȚ]\.|[A-ZĂÂÎȘȚ][a-zăâîșț]+', line):
            teacher = line; break

    return title, teacher, room

def normalize_date(s: str):
    s = (s or "").strip()
    # вычистим день недели (с/без диакритики)
    s = re.sub(r"\b(luni|marți|marti|miercuri|joi|vineri|sâmbătă|simbata|duminică|duminica)\b",
               "", s, flags=re.IGNORECASE)
    s = re.sub(r"\s+", " ", s).strip()

    m = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", s)
    if m: return s
    m = re.fullmatch(r"(\d{1,2})\.(\d{1,2})\.(\d{4})", s)
    if m:
        d, mnt, y = map(int, m.groups())
        return datetime(y, mnt, d).strftime("%Y-%m-%d")
    m = re.fullmatch(r"(\d{1,2})\.(\d{1,2})", s)
    if m:
        d, mnt = map(int, m.groups()); y = datetime.now(TZ).year
        return datetime(y, mnt, d).strftime("%Y-%m-%d")
    return None

# --- TIME ---
SUP = str.maketrans({
    "⁰":"0","¹":"1","²":"2","³":"3","⁴":"4","⁵":"5","⁶":"6","⁷":"7","⁸":"8","⁹":"9",
    "º":"0","˙":"", "’":"", "ː":":"
})
ROMAN_PREFIX = re.compile(r"^\s*(?:I|II|III|IV|V|VI|VII|VIII|IX|X)\s+", re.IGNORECASE)


# ================== DOCX parser for 4-column layout ==================
def parse_docx_to_db(path: str, target_group: str):
    doc = DocxDocument(path)
    clear_group(target_group)
    inserted = 0

    DATE_RE = re.compile(r'\b\d{1,2}\.\d{1,2}(?:\.\d{4})?\b')

    for table in doc.tables:
        nrows = len(table.rows)
        if nrows < 2:
            continue
        ncols = len(table.rows[0].cells)

        # соберём первые 2 строки как "шапку"
        hdr0 = [re.sub(r"\s+", " ", c.text).strip().lower() for c in table.rows[0].cells]
        hdr1 = [re.sub(r"\s+", " ", c.text).strip().lower() for c in table.rows[1].cells] if nrows > 1 else []
        hdr = []
        for i in range(ncols):
            h = (hdr0[i] if i < len(hdr0) else "")
            h2 = (hdr1[i] if i < len(hdr1) else "")
            if h2 and h2 not in h: h = (h + " " + h2).strip()
            hdr.append(h)

        # 1) найдём date_col, time_col по содержимому
        date_col = time_col = None
        for col in range(ncols):
            col_text = " | ".join(re.sub(r"\s+"," ", r.cells[col].text).strip() for r in table.rows[:min(8, nrows)])
            if date_col is None and ( "data" in hdr[col] or "date" in hdr[col] or DATE_RE.search(col_text) ):
                date_col = col
            if time_col is None and ( "ora" in hdr[col] or "time" in hdr[col] or parse_ora_cell(col_text)[0] ):
                time_col = col

        if date_col is None or time_col is None:
            # вероятнее всего это не нужная нам таблица
            continue

        # 2) найдём колонку нашей группы
        j_col = b_col = None
        for col in range(ncols):
            h = hdr[col]
            if "jfr" in h or "jurnalism" in h:
                j_col = col
            if "bfr" in h or "bibliotec" in h:
                b_col = col

        # если в шапке не нашли, но таблица похожа на 4-колоночную -> возьмём 2 правых колонки как группы
        if j_col is None and b_col is None and ncols >= 4:
            # предполагаем: [Data][Ora][JFR][BFR]
            j_col, b_col = (time_col + 1, time_col + 2) if time_col + 2 < ncols else (None, None)

        group_col = j_col if target_group.upper().startswith("JFR") else b_col
        if group_col is None:
            # не можем сопоставить колонку с группой — пропускаем таблицу
            continue

        # 3) найдём первую строку с данными (где есть дата/время)
        start_row = 0
        for i in range(min(6, nrows)):
            dcell = re.sub(r"\s+"," ", table.rows[i].cells[date_col].text).strip()
            tcell = re.sub(r"\s+"," ", table.rows[i].cells[time_col].text).strip()
            if normalize_date(dcell) or parse_ora_cell(tcell)[0]:
                start_row = i; break
        # подстрахуемся от шапки в 1–2 строки
        start_row = max(start_row, 1)

        last_date = None
        for r in table.rows[start_row:]:
            cells = r.cells
            if len(cells) <= max(date_col, time_col, group_col):
                continue

            # дата: падать назад на last_date, если встретился только 'joi/vineri'
            date_raw = re.sub(r"\s+"," ", (cells[date_col].text or "")).strip()
            tmp_date = normalize_date(date_raw) if date_raw else None
            date_norm = tmp_date or last_date
            if not date_norm:
                continue

            # время
            time_raw = re.sub(r"\s+"," ", (cells[time_col].text or "")).strip()
            t1, t2 = parse_ora_cell(time_raw)

            # ячейка группы
            gtxt = cells[group_col].text if group_col < len(cells) else ""
            title, teacher, room = parse_group_cell(gtxt)

            if title:
                insert_event(date_norm, t1, t2, title, teacher, room, target_group)
                inserted += 1
            last_date = date_norm

    return inserted

# ================== Import from disk ==================
def import_from_docx_path(target_group: str) -> int:
    docx_path = os.getenv("DOCX_PATH", "").strip()
    docx_glob = os.getenv("DOCX_GLOB", "").strip()

    paths = []
    if docx_path:
        p = os.path.abspath(docx_path)
        if os.path.isfile(p):
            paths = [p]
        else:
            print(f"[DOCX] DOCX_PATH указан, но файла нет: {p}")
    elif docx_glob:
        paths = sorted(glob.glob(docx_glob))
        if not paths:
            print(f"[DOCX] DOCX_GLOB указан, но файлов не найдено по маске: {docx_glob}")
    else:
        print("[DOCX] Ни DOCX_PATH, ни DOCX_GLOB не заданы в .env")

    if not paths:
        return 0

    print(f"[DOCX] Кандидаты к импорту: {paths}")
    clear_group(target_group)
    total = 0

    for p in paths:
        try:
            if not p.lower().endswith(".docx"):
                print(f"[DOCX] Пропущено (не .docx): {p}")
                continue
            if not zipfile.is_zipfile(p):
                print(f"[DOCX] Не ZIP → это невалидный .docx (возможно .doc): {p}")
                continue
            cnt = parse_docx_to_db(p, target_group=target_group)
            print(f"[DOCX] Импортировано из {os.path.basename(p)}: {cnt}")
            total += cnt
        except PackageNotFoundError:
            print(f"[DOCX] PackageNotFoundError → это не .docx: {p}")
        except Exception as e:
            print(f"[DOCX] Ошибка импорта {p}: {e}")
    return total

# ================== BOT ==================
router = Router()
USER_GROUP = {}

@router.message(Command("dump"))
async def cmd_dump(m: Message):
    from html import escape
    p = os.getenv("DOCX_PATH", "").strip()
    if not p or not os.path.isfile(p):
        return await m.answer("DOCX_PATH не задан или файл не найден.")
    doc = DocxDocument(p)
    out = [f"Таблиц: {len(doc.tables)}"]
    for ti, table in enumerate(doc.tables):
        out.append(f"\n<b>Table {ti}</b> rows={len(table.rows)} cols={len(table.rows[0].cells) if table.rows else 0}")
        for ri, row in enumerate(table.rows[:8]):  # первые 8 строк
            cells = [re.sub(r'\\s+',' ', (c.text or '').strip()) for c in row.cells]
            out.append(f"{ri:02d}: " + " | ".join(escape(x) if x else "·" for x in cells))
    await m.answer("\n".join(out))


@router.message(Command("start"))
async def cmd_start(m: Message):
    USER_GROUP[m.from_user.id] = DEFAULT_GROUP
    await m.answer(
        f"Привет! Я бот-расписание 📚\n\n"
        f"По умолчанию группа: <b>{GROUP_LABELS.get(DEFAULT_GROUP, DEFAULT_GROUP)}</b>\n"
        f"Используй кнопки ниже 👇",
        reply_markup=main_menu()
    )

@router.message(Command("help"))
async def cmd_help(m: Message):
    await m.answer(
        "Команды:\n"
        "• 📅 Сегодня — расписание на сегодня\n"
        "• 🗂 Эта неделя — все дни недели\n"
        "• 🗓 Выбрать день недели — навигация по дням\n"
        "• /reload — переимпортировать .docx из диска\n"
        "• /debug_import — показать что видим на диске\n"
        "• /docinfo — проверить файл DOCX\n"
        "• /reload_env — перечитать .env"
    )

@router.message(F.text == "📅 Сегодня")
async def today(m: Message):
    group = USER_GROUP.get(m.from_user.id, DEFAULT_GROUP)
    d = datetime.now(TZ).date()
    rows = fetch_day_for_group(d, group)
    await m.answer(fmt_day_block(d, rows, group), reply_markup=days_keyboard(d))

@router.message(F.text == "🗂 Эта неделя")
async def this_week(m: Message):
    group = USER_GROUP.get(m.from_user.id, DEFAULT_GROUP)
    start, end = week_bounds(datetime.now(TZ).date())
    rows = fetch_week_for_group(start, end, group)
    if not rows:
        return await m.answer("На эту неделю занятий нет.", reply_markup=days_keyboard(start))
    by_day = {}
    for r in rows:
        by_day.setdefault(r["date"], []).append(r)
    parts = []
    d = start
    while d <= end:
        parts.append(fmt_day_block(d, by_day.get(d.strftime("%Y-%m-%d"), []), group))
        d += timedelta(days=1)
    await m.answer("\n\n".join(parts), reply_markup=days_keyboard(start))

@router.message(F.text == "🗓 Выбрать день недели")
async def pick_day(m: Message):
    await m.answer("Выбери день:", reply_markup=days_keyboard(datetime.now(TZ).date()))

@router.callback_query(F.data.startswith("wk:"))
async def change_week(c: CallbackQuery):
    anchor = date.fromisoformat(c.data.split(":")[1])
    try:
        await safe_edit(c.message, reply_markup=days_keyboard(anchor))
    finally:
        await c.answer("Неделя обновлена")

@router.callback_query(F.data.startswith("d:"))
async def show_day(c: CallbackQuery):
    d = date.fromisoformat(c.data.split(":")[1])
    group = USER_GROUP.get(c.from_user.id, DEFAULT_GROUP)
    rows = fetch_day_for_group(d, group)
    text = fmt_day_block(d, rows, group)
    await safe_edit(c.message, text=text, reply_markup=days_keyboard(d))
    await c.answer()

@router.message(F.text == "👥 Выбрать группу")
async def choose_group(m: Message):
    await m.answer("Выбери группу:", reply_markup=groups_kb())

@router.callback_query(F.data.startswith("g:"))
async def set_group(c: CallbackQuery):
    g = c.data.split(":")[1]
    USER_GROUP[c.from_user.id] = g
    await c.answer(f"Группа установлена: {g}")
    await c.message.edit_text(f"Группа сохранена: <b>{GROUP_LABELS.get(g, g)}</b>")

@router.message(Command("reload"))
async def cmd_reload(m: Message):
    group = USER_GROUP.get(m.from_user.id, DEFAULT_GROUP)
    imported = import_from_docx_path(group)
    if imported:
        await m.answer(f"🔄 Импорт завершён.\nГруппа: <b>{group}</b>\nЗаписей: <b>{imported}</b>")
    else:
        await m.answer("Файл(ы) не найдены или парсер не нашёл пар. Проверь .env / формат.")

@router.message(Command("debug_import"))
async def debug_import(m: Message):
    here = os.getcwd()
    docx_path = os.getenv("DOCX_PATH", "")
    docx_glob = os.getenv("DOCX_GLOB", "")
    files_here = sorted([f for f in os.listdir(here) if f.lower().endswith((".docx", ".doc"))])
    globbed = sorted(glob.glob(docx_glob)) if docx_glob else []
    msg = (
        "🛠 Debug import\n"
        f"cwd: <code>{here}</code>\n"
        f"DOCX_PATH: <code>{docx_path}</code>\n"
        f"DOCX_GLOB: <code>{docx_glob}</code>\n"
        f"В папке (.doc/.docx): {files_here}\n"
        f"По маске: {globbed}"
    )
    await m.answer(msg)

@router.message(Command("docinfo"))
async def cmd_docinfo(m: Message):
    p = os.getenv("DOCX_PATH", "").strip()
    if not p:
        return await m.answer("DOCX_PATH не задан в .env")
    ap = os.path.abspath(p)
    if not os.path.isfile(ap):
        return await m.answer(f"Файл не найден:\n<code>{ap}</code>")
    try:
        info = [f"Путь: <code>{ap}</code>", f"Размер: {os.path.getsize(ap)} байт", f"ZIP: {zipfile.is_zipfile(ap)}"]
        doc = DocxDocument(ap)
        info += [f"Параграфов: {len(doc.paragraphs)}", f"Таблиц: {len(doc.tables)}"]
        await m.answer("🧪 DOCX info:\n" + "\n".join(info))
    except PackageNotFoundError:
        await m.answer("❌ Это не .docx (или файл битый). Конвертируй .doc → .docx.")
    except Exception as e:
        await m.answer(f"❌ Не смог открыть через python-docx:\n<code>{e}</code>")

@router.message(Command("reload_env"))
async def reload_env(m: Message):
    load_dotenv(override=True)
    await m.answer("♻️ .env перечитан. Запусти /reload для повторного импорта.")

# ================== MAIN ==================
async def main():
    init_db()
    imported = import_from_docx_path(DEFAULT_GROUP)
    if imported:
        print(f"[DOCX] Авто-импорт: {imported} записей для {DEFAULT_GROUP}")
    else:
        print("[DOCX] Авто-импорт: файлов не найдено или пар не извлечено")

    bot = Bot(BOT_TOKEN, default=DefaultBotProperties(parse_mode="HTML"))
    dp = Dispatcher()
    dp.include_router(router)
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
